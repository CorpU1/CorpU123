import sys
from PyQt5 import QtCore, QtGui, QtWidgets
from docx import Document
import os

class Ui_MainWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(812, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.formLayout = QtWidgets.QFormLayout(self.centralwidget)
        self.formLayout.setObjectName("formLayout")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setObjectName("label_3")
        self.formLayout.setWidget(0, QtWidgets.QFormLayout.LabelRole, self.label_3)
        self.textEdit_2 = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_2.setObjectName("textEdit_2")
        self.formLayout.setWidget(0, QtWidgets.QFormLayout.FieldRole, self.textEdit_2)
        self.label_6 = QtWidgets.QLabel(self.centralwidget)
        self.label_6.setObjectName("label_6")
        self.formLayout.setWidget(2, QtWidgets.QFormLayout.LabelRole, self.label_6)
        self.textEdit_3 = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_3.setObjectName("textEdit_3")
        self.formLayout.setWidget(2, QtWidgets.QFormLayout.FieldRole, self.textEdit_3)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setObjectName("label")
        self.formLayout.setWidget(4, QtWidgets.QFormLayout.LabelRole, self.label)
        self.dateEdit = QtWidgets.QDateEdit(self.centralwidget)
        self.dateEdit.setObjectName("dateEdit")
        self.formLayout.setWidget(4, QtWidgets.QFormLayout.FieldRole, self.dateEdit)
        self.label_7 = QtWidgets.QLabel(self.centralwidget)
        self.label_7.setObjectName("label_7")
        self.formLayout.setWidget(6, QtWidgets.QFormLayout.LabelRole, self.label_7)
        self.textEdit_6 = QtWidgets.QTextEdit(self.centralwidget)
        self.textEdit_6.setObjectName("textEdit_6")
        self.formLayout.setWidget(6, QtWidgets.QFormLayout.FieldRole, self.textEdit_6)
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setObjectName("label_5")
        self.formLayout.setWidget(7, QtWidgets.QFormLayout.LabelRole, self.label_5)
        self.spinBox_2 = QtWidgets.QSpinBox(self.centralwidget)
        self.spinBox_2.setObjectName("spinBox_2")
        self.formLayout.setWidget(7, QtWidgets.QFormLayout.FieldRole, self.spinBox_2)
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setObjectName("label_4")
        self.formLayout.setWidget(10, QtWidgets.QFormLayout.LabelRole, self.label_4)
        self.spinBox = QtWidgets.QSpinBox(self.centralwidget)
        self.spinBox.setObjectName("spinBox")
        self.formLayout.setWidget(10, QtWidgets.QFormLayout.FieldRole, self.spinBox)
        self.horizontalLayout = QtWidgets.QHBoxLayout()
        self.horizontalLayout.setObjectName("horizontalLayout")
        spacerItem = QtWidgets.QSpacerItem(40, 20, QtWidgets.QSizePolicy.Expanding, QtWidgets.QSizePolicy.Minimum)
        self.horizontalLayout.addItem(spacerItem)
        self.pushButton_2 = QtWidgets.QPushButton(self.centralwidget)
        self.pushButton_2.setObjectName("pushButton_2")
        self.horizontalLayout.addWidget(self.pushButton_2)
        self.formLayout.setLayout(12, QtWidgets.QFormLayout.SpanningRole, self.horizontalLayout)
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 812, 18))
        self.menubar.setObjectName("menubar")
        self.menu = QtWidgets.QMenu(self.menubar)
        self.menu.setObjectName("menu")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.menubar.addAction(self.menu.menuAction())

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label_3.setText(_translate("MainWindow", "Цели"))
        self.label_6.setText(_translate("MainWindow", "Описание бизнеса"))
        self.label.setText(_translate("MainWindow", "Дата"))
        self.label_7.setText(_translate("MainWindow", "Действия, принятые на первом  ОСУ"))
        self.label_5.setText(_translate("MainWindow", "Доля Участника 1"))
        self.label_4.setText(_translate("MainWindow", "Количество СД"))
        self.pushButton_2.setText(_translate("MainWindow", "ОК"))
        self.menu.setTitle(_translate("MainWindow", "КорпСтудия"))


from PyQt5.QtCore import QDate


class CorporateApp(QtWidgets.QMainWindow):
    def __init__(self):
        super().__init__()
        self.ui = Ui_MainWindow()
        self.ui.setupUi(self)
        self.ui.pushButton_2.clicked.connect(self.generate_contract)

    def generate_contract(self):
        sd = self.ui.spinBox.value()
        stock = self.ui.spinBox_2.value()
        stucke = 100 - stock
        decisions = self.ui.textEdit_6.toPlainText()
        business = self.ui.textEdit_3.toPlainText()
        target = self.ui.textEdit_2.toPlainText()

        date = self.ui.dateEdit.date()
        formatted_date = date.toString("dd.MM.yyyy")

        print(f"sd: {sd}, stock: {stock}, stucke: {stucke}")
        print(f"decisions: {decisions}")
        print(f"business: {business}")
        print(f"target: {target}")
        print(f"formatted_date: {formatted_date}")

        options = QtWidgets.QFileDialog.Options()
        file_path, _ = QtWidgets.QFileDialog.getSaveFileName(self, "Сохранить договор", "",
                                                             "Word Documents (*.docx);;All Files (*)", options=options)

        if file_path:
            try:
                self.create_contract(sd, stock, stucke, decisions, business, target, formatted_date, file_path)
                QtWidgets.QMessageBox.information(self, "Успех", "Договор успешно создан!")
            except Exception as e:
                QtWidgets.QMessageBox.critical(self, "Ошибка", f"Не удалось создать договор: {str(e)}")

    def create_contract(self, sd, stock, stucke, decisions, business, target, formatted_date, output_path):
        template_path = 'C:/oldqt/SavedCorpDoc/Corporate1.docx'
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Шаблон договора не найден: {template_path}")

        doc = Document(template_path)

        def replace_text_in_runs(paragraph, placeholder, replacement):
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)
                elif placeholder.lower() in run.text.lower():
                    run.text = run.text.replace(placeholder, replacement)

        try:
            for p in doc.paragraphs:
                replace_text_in_runs(p, 'SD', str(sd))
                replace_text_in_runs(p, 'stock', str(stock))
                replace_text_in_runs(p, 'stucke', str(stucke))
                replace_text_in_runs(p, 'decisions', decisions)
                replace_text_in_runs(p, 'business', business)
                replace_text_in_runs(p, 'target', target)
                replace_text_in_runs(p, 'date', formatted_date)  # Замена даты

            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            replace_text_in_runs(p, 'SD', str(sd))
                            replace_text_in_runs(p, 'stock', str(stock))
                            replace_text_in_runs(p, 'stucke', str(stucke))
                            replace_text_in_runs(p, 'decisions', decisions)
                            replace_text_in_runs(p, 'business', business)
                            replace_text_in_runs(p, 'target', target)
                            replace_text_in_runs(p, 'date', formatted_date)  # Замена даты

            doc.save(output_path)
            print(f"Document saved at {output_path}")
            QtWidgets.QMessageBox.information(self, "Успех", "Договор успешно создан!")
        except Exception as e:
            QtWidgets.QMessageBox.critical(self, "Ошибка", f"Не удалось создать договор: {str(e)}")


if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    window = CorporateApp()
    window.show()
    sys.exit(app.exec_())
